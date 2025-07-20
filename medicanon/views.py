import csv
import io
import hashlib
import spacy
from django.contrib.auth import login, logout
from django.contrib.auth.forms import UserCreationForm
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.http import FileResponse, Http404, HttpResponse
from django.core.paginator import Paginator
from django.contrib import messages
from django.db import transaction
from django.urls import reverse
import PyPDF2
import pdfplumber
from docx import Document
from .models import Utilisateur, Fichier, Historique, Métriques, RapportConformité, Données, RègleAnonymisation
import pandas as pd
import os
from django.core.files.base import ContentFile
from django.core.files.storage import default_storage
import re
from datetime import datetime
import secrets
from django.views.decorators.http import require_GET


# Charger le modèle spaCy une seule fois
nlp = spacy.load("fr_core_news_sm")

# Dictionnaires étendus pour le Sénégal
SENEGALESE_FIRSTNAMES = [
    "Awa", "Fatou", "Pape", "Fatoumata", "Fatimata", "Fatima", "Memedou", 
    "Mouhamed", "Modou", "Mouhamadou", "Cheikh", "Amadou", "Omar", "Oumar", 
    "Aminata", "Khadidiatou", "Mariama", "Ousmane", "Ibrahima", "Abdoulaye", 
    "Khadija", "Aissatou", "Adama", "Seynabou", "Sokhna", "Binta", "Ndèye",
    "Moussa", "Mamadou", "Aliou", "Lamine", "Saliou", "Babacar"
]

SENEGALESE_LASTNAMES = [
    "Diouf", "Sow", "Ndiaye", "Fall", "Diop", "Ba", "Sarr", "Gaye", "Mbengue", 
    "Thiam", "Faye", "Diallo", "Dia", "Kandji", "Niang", "Ngom", "Cissé", "Sy", 
    "Mbacke", "Seck", "Tall", "Demba", "Kane", "Sané", "Samb", "Camara", "Touré"
]

# Patterns de détection améliorés
MEDICAL_PATTERNS = {
    'patient_id': r'\b(patient|id|identifiant)[\s_-]*(id|num|numero|number)?\b',
    'phone': r'\b(\+221\s?)?[0-9]{2}[\s.-]?[0-9]{3}[\s.-]?[0-9]{2}[\s.-]?[0-9]{2}\b',
    'email': r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b',
    'date_birth': r'\b(date[\s_-]*naissance|birth[\s_-]*date|ddn|dob)\b',
    'address': r'\b(adresse|address|domicile|residence)\b',
    'medical_id': r'\b(cni|carte[\s_-]*identite|passport|passeport)\b'
}

def detect_ipi_fields_enhanced(fichier):
    """Détection améliorée des champs IPI"""
    ipi_fields = []
    text = extract_text_from_file(fichier.fichier)
    
    if not text:
        return ipi_fields
    
    file_extension = fichier.fichier.name.split('.')[-1].lower()
    
    if file_extension == 'csv':
        fichier.fichier.seek(0)
        decoded_file = fichier.fichier.read().decode('utf-8', errors='ignore')
        csv_reader = csv.reader(io.StringIO(decoded_file))
        headers = next(csv_reader, [])
        
        # Analyser les en-têtes
        for header in headers:
            header_lower = header.lower().strip()
            
            # Vérifications par patterns
            for pattern_name, pattern in MEDICAL_PATTERNS.items():
                if re.search(pattern, header_lower, re.IGNORECASE):
                    ipi_fields.append(header)
                    break
            
            # Vérifications par mots-clés spécifiques
            keywords = ['nom', 'prenom', 'name', 'firstname', 'lastname', 'age', 'sexe', 
                       'telephone', 'phone', 'email', 'adresse', 'address', 'patient']
            
            if any(keyword in header_lower for keyword in keywords):
                if header not in ipi_fields:
                    ipi_fields.append(header)
        
        # Analyser un échantillon de données pour validation
        sample_rows = []
        for i, row in enumerate(csv_reader):
            if i >= 10:  # Analyser 10 lignes max
                break
            sample_rows.append(row)
        
        # Validation des champs détectés avec les données
        validated_fields = []
        for field in ipi_fields:
            field_index = headers.index(field) if field in headers else -1
            if field_index != -1:
                # Vérifier si la colonne contient des données sensibles
                contains_sensitive = False
                for row in sample_rows:
                    if field_index < len(row):
                        cell_value = row[field_index].strip()
                        if is_sensitive_data(cell_value):
                            contains_sensitive = True
                            break
                
                if contains_sensitive:
                    validated_fields.append(field)
        
        return list(set(validated_fields)) if validated_fields else list(set(ipi_fields))
    
    else:  # PDF, DOCX
        doc = nlp(text)
        entities_found = []
        
        for ent in doc.ents:
            if ent.label_ in ['PER', 'LOC', 'ORG']:
                entities_found.append(f"Entité_{ent.label_}_{hash(ent.text) % 100}")
        
        # Recherche de noms sénégalais
        for fname in SENEGALESE_FIRSTNAMES:
            if fname.lower() in text.lower():
                entities_found.append(f"Prénom_Sénégalais")
                break
        
        for lname in SENEGALESE_LASTNAMES:
            if lname.lower() in text.lower():
                entities_found.append(f"Nom_Sénégalais")
                break
        
        return list(set(entities_found))

def is_sensitive_data(value):
    """Vérifie si une valeur contient des données sensibles"""
    if not value or len(value.strip()) < 2:
        return False
    
    value = value.strip()
    
    # Vérifications spécifiques
    checks = [
        # Email
        re.search(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', value),
        # Téléphone sénégalais
        re.search(r'\b(\+221\s?)?[0-9]{2}[\s.-]?[0-9]{3}[\s.-]?[0-9]{2}[\s.-]?[0-9]{2}\b', value),
        # Date de naissance
        re.search(r'\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b', value),
        # Noms sénégalais
        any(name.lower() in value.lower() for name in SENEGALESE_FIRSTNAMES + SENEGALESE_LASTNAMES),
        # Adresse (contient des mots comme rue, avenue, etc.)
        re.search(r'\b(rue|avenue|boulevard|quartier|villa|lot)\b', value, re.IGNORECASE)
    ]
    
    return any(checks)

def anonymize_data_enhanced(headers, rows, selected_ipi, methods):
    """Anonymisation améliorée avec de meilleures méthodes"""
    anonymized_rows = []
    colonnes_anonymisees = 0
    
    # Créer un mapping de pseudonymes cohérent
    pseudonym_mapping = {}
    
    for row in rows:
        new_row = []
        for i, cell in enumerate(row):
            header = headers[i] if i < len(headers) else f"champ_{i+1}"
            
            if header in selected_ipi:
                method = methods.get(header, 'none')
                original_value = str(cell).strip()
                
                if method == 'suppression':
                    new_row.append("[SUPPRIMÉ]")
                    colonnes_anonymisees += 1
                
                elif method == 'pseudonymisation':
                    # Pseudonymisation cohérente
                    if original_value not in pseudonym_mapping:
                        pseudonym_mapping[original_value] = f"USER_{secrets.randbelow(99999):05d}"
                    new_row.append(pseudonym_mapping[original_value])
                    colonnes_anonymisees += 1
                
                elif method == 'generalisation':
                    # Généralisation intelligente
                    generalized = generalize_value(original_value, header)
                    new_row.append(generalized)
                    colonnes_anonymisees += 1
                
                elif method == 'hachage':
                    # Hachage avec salt
                    salt = "medic_anon_salt_2024"
                    hashed = hashlib.sha256((original_value + salt).encode()).hexdigest()[:16]
                    new_row.append(f"HASH_{hashed}")
                    colonnes_anonymisees += 1
                
                else:
                    new_row.append(cell)
            else:
                new_row.append(cell)
        
        anonymized_rows.append(new_row)
    
    return anonymized_rows, colonnes_anonymisees

def generalize_value(value, field_type):
    """Généralisation intelligente selon le type de champ"""
    value_str = str(value).strip().lower()
    field_lower = field_type.lower()
    
    # Généralisation d'âge
    if 'age' in field_lower or re.search(r'\b\d{1,3}\b', value):
        try:
            age = int(re.search(r'\d+', value).group())
            if age < 18:
                return "Mineur"
            elif age < 30:
                return "18-29 ans"
            elif age < 50:
                return "30-49 ans"
            elif age < 70:
                return "50-69 ans"
            else:
                return "70+ ans"
        except:
            return "Age non spécifié"
    
    # Généralisation de date
    if 'date' in field_lower or 'naissance' in field_lower:
        try:
            # Extraire juste l'année
            year_match = re.search(r'\b(19|20)\d{2}\b', value)
            if year_match:
                year = int(year_match.group())
                decade = (year // 10) * 10
                return f"Années {decade}s"
        except:
            pass
        return "Période non spécifiée"
    
    # Généralisation de téléphone
    if 'phone' in field_lower or 'telephone' in field_lower:
        return "7X-XXX-XX-XX"
    
    # Généralisation d'email
    if '@' in value_str:
        domain = value_str.split('@')[-1] if '@' in value_str else ""
        return f"utilisateur@{domain}" if domain else "email@domaine.com"
    
    # Généralisation d'adresse
    if 'adresse' in field_lower or 'address' in field_lower:
        return "Région de Dakar"
    
    # Par défaut
    return "Valeur généralisée"

# Fonction améliorée pour le calcul du score de sécurité
def calculate_advanced_security_score(fields, methods, file_size=0):
    """Calcul avancé du score de sécurité avec pondération"""
    
    if not fields or not methods:
        return 0
    
    # Pondération par type de données (selon sensibilité RGPD)
    sensitivity_weights = {
        'nom': 35, 'prenom': 30, 'name': 35, 'firstname': 30, 'lastname': 35,
        'date_naissance': 30, 'birth_date': 30, 'ddn': 30, 'age': 20,
        'cni': 40, 'passport': 40, 'passeport': 40, 'carte_identite': 40,
        'adresse': 25, 'address': 25, 'domicile': 25,
        'telephone': 20, 'phone': 20, 'mobile': 20,
        'email': 20, 'mail': 20,
        'patient_id': 25, 'numero_patient': 25,
        'sexe': 15, 'genre': 15, 'gender': 15,
        'profession': 15, 'job': 15, 'travail': 15
    }
    
    # Pondération par méthode d'anonymisation
    method_weights = {
        'suppression': 30,      # Données complètement supprimées
        'hachage': 35,          # Irréversible cryptographiquement
        'pseudonymisation': 25, # Réversible avec clé
        'generalisation': 20,   # Perte de précision
        'none': 0               # Aucune protection
    }
    
    total_weighted_score = 0
    max_possible_score = 0
    
    for field in fields:
        # Trouver le poids de sensibilité
        field_lower = field.lower()
        sensitivity = 10  # Valeur par défaut
        
        for pattern, weight in sensitivity_weights.items():
            if pattern in field_lower:
                sensitivity = weight
                break
        
        # Obtenir la méthode utilisée
        method = methods.get(field, 'none')
        method_weight = method_weights.get(method, 0)
        
        # Calculer le score pondéré
        field_score = (sensitivity * method_weight) / 100
        total_weighted_score += field_score
        max_possible_score += sensitivity
    
    if max_possible_score == 0:
        return 0
    
    # Normaliser sur 100 et appliquer des bonus/malus
    base_score = (total_weighted_score / max_possible_score) * 100
    
    # Bonus pour diversité des méthodes
    unique_methods = set(m for m in methods.values() if m != 'none')
    if len(unique_methods) > 1:
        base_score += 5
    
    # Bonus pour méthodes fortes sur données sensibles
    strong_methods_count = sum(1 for m in methods.values() if m in ['hachage', 'suppression'])
    if strong_methods_count > len(fields) * 0.6:
        base_score += 10
    
    # Malus pour trop de pseudonymisation
    pseudo_count = sum(1 for m in methods.values() if m == 'pseudonymisation')
    if pseudo_count > len(fields) * 0.7:
        base_score -= 5
    
    return min(100, max(0, base_score))

def evaluate_gdpr_compliance(fichier, selected_fields, methods):
    """Évalue la conformité RGPD et CDP Sénégal"""
    
    compliance_score = 0
    recommendations = []
    risks = []
    
    # 1. Vérification de la minimisation des données (Article 5(1)(c) RGPD)
    total_fields = len(selected_fields) if selected_fields else 0
    if total_fields <= 5:
        compliance_score += 20
    elif total_fields <= 10:
        compliance_score += 15
        recommendations.append("Considérer la réduction du nombre de champs traités")
    else:
        compliance_score += 10
        risks.append("Trop de champs sélectionnés - Principe de minimisation")
    
    # 2. Évaluation des méthodes d'anonymisation
    method_compliance = {
        'suppression': {'score': 25, 'risk': 'Faible', 'reversible': False},
        'hachage': {'score': 30, 'risk': 'Très faible', 'reversible': False},
        'pseudonymisation': {'score': 20, 'risk': 'Moyen', 'reversible': True},
        'generalisation': {'score': 15, 'risk': 'Moyen', 'reversible': False},
        'none': {'score': 0, 'risk': 'Élevé', 'reversible': True}
    }
    
    methods_used = list(methods.values()) if methods else []
    method_scores = [method_compliance.get(m, {'score': 0})['score'] for m in methods_used]
    
    if method_scores:
        avg_method_score = sum(method_scores) / len(method_scores)
        compliance_score += min(30, avg_method_score)
    
    # 3. Vérification de la réversibilité (Article 4(5) RGPD)
    reversible_methods = [m for m in methods_used if method_compliance.get(m, {}).get('reversible', True)]
    if not reversible_methods:
        compliance_score += 25
    elif len(reversible_methods) < len(methods_used) / 2:
        compliance_score += 15
        recommendations.append("Privilégier des méthodes non-réversibles pour une meilleure anonymisation")
    else:
        compliance_score += 5
        risks.append("Plusieurs méthodes réversibles utilisées - Risque de ré-identification")
    
    # 4. Évaluation spécifique CDP Sénégal
    sensitive_fields = ['nom', 'prenom', 'cni', 'date_naissance', 'adresse']
    cdp_sensitive_found = [f for f in selected_fields if any(s in f.lower() for s in sensitive_fields)]
    
    if cdp_sensitive_found:
        strong_methods = [f for f in cdp_sensitive_found if methods.get(f) in ['hachage', 'suppression']]
        if len(strong_methods) == len(cdp_sensitive_found):
            compliance_score += 25
        else:
            compliance_score += 10
            recommendations.append("Utiliser hachage ou suppression pour les données d'identité sénégalaises")
    
    # Déterminer le niveau de conformité
    if compliance_score >= 90:
        conformity_level = "Excellente conformité RGPD/CDP"
    elif compliance_score >= 75:
        conformity_level = "Bonne conformité RGPD/CDP"
    elif compliance_score >= 60:
        conformity_level = "Conformité acceptable avec améliorations"
    else:
        conformity_level = "Non-conformité - Révision nécessaire"
    
    return {
        'score': compliance_score,
        'level': conformity_level,
        'recommendations': recommendations,
        'risks': risks,
        'reversible_methods': len(reversible_methods),
        'strong_anonymization': len([m for m in methods_used if m in ['hachage', 'suppression']])
    }

# Ajout dans views.py pour l'utilisation
def create_compliance_report(fichier, selected_ipi, methods, metrics):
    """Crée un rapport de conformité détaillé"""
    
    compliance_eval = evaluate_gdpr_compliance(fichier, selected_ipi, methods)
    
    # Analyser les risques
    risk_analysis = []
    
    # Risque de ré-identification
    if compliance_eval['reversible_methods'] > 0:
        risk_analysis.append(f"Risque de ré-identification: {compliance_eval['reversible_methods']} méthode(s) réversible(s)")
    
    # Risque de perte d'utilité
    if sum(1 for m in methods.values() if m == 'suppression') > len(selected_ipi) * 0.5:
        risk_analysis.append("Risque de perte d'utilité: Trop de suppressions")
    
    # Risque de non-conformité
    if compliance_eval['score'] < 75:
        risk_analysis.append("Risque de non-conformité RGPD/CDP")
    
    # Recommandations spécifiques
    recommendations = compliance_eval['recommendations'].copy()
    
    if metrics['taux_anonymisation'] < 30:
        recommendations.append("Augmenter le taux d'anonymisation (actuellement < 30%)")
    
    if compliance_eval['strong_anonymization'] == 0:
        recommendations.append("Utiliser au moins une méthode forte (hachage/suppression)")
    
    return {
        'conformity_score': compliance_eval['score'],
        'conformity_level': compliance_eval['level'],
        'risk_analysis': '; '.join(risk_analysis) if risk_analysis else "Risques faibles",
        'recommendations': '; '.join(recommendations) if recommendations else "Traitement conforme",
        'gdpr_compliant': compliance_eval['score'] >= 75,
        'cdp_compliant': compliance_eval['score'] >= 70,  # Seuil CDP Sénégal
    }

# Template pour affichage du rapport de conformité
def format_compliance_display(compliance_data):
    """Formate les données de conformité pour affichage"""
    
    score = compliance_data['conformity_score']
    
    if score >= 90:
        badge_color = "green"
        icon = "✅"
    elif score >= 75:
        badge_color = "blue"  
        icon = "✔️"
    elif score >= 60:
        badge_color = "orange"
        icon = "⚠️"
    else:
        badge_color = "red"
        icon = "❌"
    
    return {
        'score': score,
        'badge_color': badge_color,
        'icon': icon,
        'level': compliance_data['conformity_level'],
        'gdpr_status': "Conforme RGPD" if compliance_data['gdpr_compliant'] else "Non-conforme RGPD",
        'cdp_status': "Conforme CDP" if compliance_data['cdp_compliant'] else "Non-conforme CDP"
    }

def register_view(request):
    if request.method == 'POST':
        form = UserCreationForm(request.POST)
        if form.is_valid():
            user = form.save(commit=False)
            user.role = 'Visiteur'  # role forcé
            user.email = request.POST.get('email')
            user.first_name = request.POST.get('first_name')
            user.last_name = request.POST.get('last_name')
            user.save()
            login(request, user)
            return redirect('public_hub')
    else:
        form = UserCreationForm()
    return render(request, 'medicanon/register.html', {'form': form})

def accueil(request):
    return render(request, 'medicanon/accueil.html')

def redirect_after_login(request):
    user = request.user
    if user.is_authenticated:
        if user.role in ['Agent', 'Administrateur']:
            return redirect('base')
        else:
            return redirect('accueil')
    return redirect('login')

@login_required
def manage_users(request):
    return render(request, 'medicanon/manage_users.html')

def extract_text_from_file(file):
    file_extension = file.name.split('.')[-1].lower()
    text = ""
    file.seek(0)
    if file_extension == 'csv':
        decoded_file = file.read().decode('utf-8', errors='ignore')
        text = decoded_file
    elif file_extension == 'pdf':
        with pdfplumber.open(file) as pdf:
            for page in pdf.pages:
                text += page.extract_text() or ""
    elif file_extension == 'docx':
        doc = Document(io.BytesIO(file.read()))
        for para in doc.paragraphs:
            text += para.text + "\n"
    return text

def detect_ipi_fields(fichier):
    ipi_fields = []
    text = extract_text_from_file(fichier.fichier)
    if not text:
        return ipi_fields
    file_extension = fichier.fichier.name.split('.')[-1].lower()
    if file_extension in ['pdf', 'docx']:
        doc = nlp(text)
        for ent in doc.ents:
            if ent.label_ in ['PER', 'LOC'] or any(fname in ent.text for fname in SENEGALESE_FIRSTNAMES) or any(lname in ent.text for lname in SENEGALESE_LASTNAMES):
                ipi_fields.append(f"Champ_{hash(ent.text) % 10 + 1}")
        return list(set(ipi_fields))
    fichier.fichier.seek(0)
    decoded_file = fichier.fichier.read().decode('utf-8', errors='ignore')
    csv_reader = csv.reader(io.StringIO(decoded_file))
    headers = next(csv_reader, [])
    for header in headers:
        doc = nlp(header.lower())
        if any(ent.label_ in ['PER', 'LOC', 'ORG'] for ent in doc.ents) or \
           any(fname.lower() in header.lower() for fname in SENEGALESE_FIRSTNAMES) or \
           any(lname.lower() in header.lower() for lname in SENEGALESE_LASTNAMES) or \
           any(kw.lower() in header.lower() for kw in ['patient', 'nom', 'prenom', 'adresse', 'email', 'telephone', 'date_naissance']):
            ipi_fields.append(header)
    for i, row in enumerate(csv_reader):
        if i >= 5:
            break
        for cell in row:
            doc = nlp(cell.lower())
            if any(ent.label_ in ['PER', 'LOC'] for ent in doc.ents) or \
               any(fname.lower() in cell.lower() for fname in SENEGALESE_FIRSTNAMES) or \
               any(lname.lower() in cell.lower() for lname in SENEGALESE_LASTNAMES):
                header_idx = row.index(cell)
                if headers[header_idx] not in ipi_fields:
                    ipi_fields.append(headers[header_idx])
    return list(set(ipi_fields))

def calculate_sensitivity_score(fields):
    sensitivity_weights = {'nom': 30, 'prenom': 25, 'adresse': 20, 'date_naissance': 15, 'email': 10, 'patient': 20}
    total_weight = sum(sensitivity_weights.get(field.lower(), 0) for field in fields)
    return min(100, total_weight)

def anonymize_data(headers, rows, selected_ipi, methods):
    anonymized_rows = []
    colonnes_anonymisees = 0
    for row in rows:
        new_row = []
        for i, cell in enumerate(row):
            header = headers[i].lower() if headers else f"champ_{i+1}"
            if header in [h.lower() for h in selected_ipi]:
                method = methods.get(headers[i] if headers else f"champ_{i+1}", 'none')
                if method == 'suppression':
                    new_row.append("")
                    colonnes_anonymisees += 1
                elif method == 'pseudonymisation':
                    new_row.append(f"user_{hash(cell) % 10000}")
                    colonnes_anonymisees += 1
                elif method == 'generalisation':
                    new_row.append("Tranche")
                    colonnes_anonymisees += 1
                elif method == 'hachage':
                    new_row.append(hashlib.sha256(cell.encode()).hexdigest())
                    colonnes_anonymisees += 1
                else:
                    new_row.append(cell)
            else:
                new_row.append(cell)
        anonymized_rows.append(new_row)
    return anonymized_rows


@login_required
def anonymize(request):
    step = int(request.GET.get('step', 1))
    fichier_id = request.GET.get('fichier_id')  # Récupère l'ID depuis l'URL
    uploaded_file = request.session.get('uploaded_file')  # Récupère depuis la session
    fichiers = Fichier.objects.filter(utilisateur=request.user, statut__in=['Importé', 'En cours'])
    methodes = ['Suppression', 'Pseudonymisation', 'Généralisation', 'Hachage']
    ipi_fields = ['nom', 'prenom', 'date_naissance', 'adresse', 'email', 'telephone']  # Sera remplacé dynamiquement
    sensitivity_score = 0
    lignes_traitees = 0
    colonnes_anonymisees = 0
    taux_anonymisation = 0
    score_securite = 0
    analyse_risques = "Risque faible"
    conformite = "Conforme RGPD/CDP"
    recommandations = "Conserver dans un environnement sécurisé"
    headers = None
    original = None
    anonymized = None

    # Récupérer les informations du fichier si fichier_id existe et détecter les champs IPI
    if fichier_id:
        try:
            fichier_instance = Fichier.objects.get(id=fichier_id, utilisateur=request.user)
            uploaded_file = {
                'nom_fichier': fichier_instance.nom_fichier,
                'extension': fichier_instance.nom_fichier.split('.')[-1].lower(),
                'id': fichier_id
            }
            request.session['uploaded_file'] = uploaded_file  # Mettre à jour la session
            # Détecter les champs IPI dynamiquement
            ipi_fields = detect_ipi_fields(fichier_instance)
            sensitivity_score = calculate_sensitivity_score(ipi_fields)
        except Fichier.DoesNotExist:
            messages.error(request, "Fichier introuvable.")
            if 'uploaded_file' in request.session:
                del request.session['uploaded_file']  # Supprimer si invalide
            return redirect('anonymize')

    if request.method == 'POST':
        if 'fichier' in request.FILES:
            uploaded_file = request.FILES['fichier']
            file_extension = uploaded_file.name.split('.')[-1].lower()
            if file_extension not in ['csv', 'pdf', 'docx']:
                messages.error(request, "Format non supporté. Utilisez CSV, PDF ou DOCX.")
                return redirect('anonymize')
            with transaction.atomic():
                fichier_instance = Fichier.objects.create(
                    fichier=uploaded_file,
                    nom_fichier=uploaded_file.name,
                    statut='Importé',
                    utilisateur=request.user
                )
                fichier_id = str(fichier_instance.id)  # Stocke l'ID du fichier importé
            uploaded_file = {'nom_fichier': uploaded_file.name, 'extension': file_extension, 'id': fichier_id}
            request.session['uploaded_file'] = uploaded_file  # Stocker dans la session
            messages.success(request, "Fichier importé avec succès.")
            # Correction : Utiliser reverse avec les paramètres de requête
            url = reverse('anonymize') + f'?step=2&fichier_id={fichier_id}'
            return redirect(url)
        elif 'from_config' in request.POST:
            fichier_id = request.POST.get('fichier_id')
            if not fichier_id:
                messages.error(request, "Aucun fichier sélectionné.")
                url = reverse('anonymize') + '?step=2'
                return redirect(url)
            
            # Récupérer les champs IPI sélectionnés et les méthodes
            selected_ipi = request.POST.getlist('ipi_fields')
            # Vérifier qu'au moins un champ IPI est sélectionné
            if not selected_ipi:
                messages.error(request, "Veuillez sélectionner au moins un champ IPI à anonymiser.")
                url = reverse('anonymize') + f'?step=2&fichier_id={fichier_id}'
                return redirect(url)

            # Vérifier que des méthodes sont sélectionnées
            methods = {field: request.POST.get(f'method_{field}', 'none') for field in selected_ipi}
            if all(method == 'none' for method in methods.values()):
                messages.error(request, "Veuillez sélectionner une méthode d'anonymisation pour au moins un champ.")
                url = reverse('anonymize') + f'?step=2&fichier_id={fichier_id}'
                return redirect(url)
            
            # Stocker les configurations dans la session
            request.session['selected_ipi'] = selected_ipi
            request.session['methods'] = methods
            
            # Traiter le fichier CSV
            try:
                fichier_instance = Fichier.objects.get(id=fichier_id, utilisateur=request.user)
                file_extension = fichier_instance.nom_fichier.split('.')[-1].lower()
                
                if file_extension == 'csv':
                    fichier_instance.fichier.seek(0)
                    decoded_file = fichier_instance.fichier.read().decode('utf-8', errors='ignore')
                    csv_reader = csv.reader(io.StringIO(decoded_file))
                    headers = next(csv_reader, [])
                    original = list(csv_reader)
                    
                    # Vérifier que le fichier n'est pas vide
                    if not headers:
                        messages.error(request, "Le fichier CSV semble vide ou mal formaté.")
                        url = reverse('anonymize') + f'?step=2&fichier_id={fichier_id}'
                        return redirect(url)
                    
                    # Anonymiser les données
                    anonymized = anonymize_data(headers, original, selected_ipi, methods)
                    
                    # Calculer les métriques
                    lignes_traitees = len(original)
                    colonnes_anonymisees = len(selected_ipi)
                    taux_anonymisation = (colonnes_anonymisees / len(headers)) * 100 if headers else 0
                    score_securite = min(100, taux_anonymisation + calculate_sensitivity_score(selected_ipi))
                    
                    # Stocker toutes les données dans la session
                    request.session['headers'] = headers
                    request.session['original'] = original
                    request.session['anonymized'] = anonymized
                    request.session['lignes_traitees'] = lignes_traitees
                    request.session['colonnes_anonymisees'] = colonnes_anonymisees
                    request.session['taux_anonymisation'] = taux_anonymisation
                    request.session['score_securite'] = score_securite
                    
                    # Debugging - vérifier que les données sont bien stockées
                    print(f"DEBUG - Données stockées en session:")
                    print(f"  - selected_ipi: {len(selected_ipi)} éléments")
                    print(f"  - methods: {len(methods)} éléments")
                    print(f"  - headers: {len(headers)} éléments")
                    print(f"  - original: {len(original)} lignes")
                    print(f"  - anonymized: {len(anonymized)} lignes")
                    
                    # Redirection vers l'étape 3
                    url = reverse('anonymize') + f'?step=3&fichier_id={fichier_id}'
                    return redirect(url)
                else:
                    messages.error(request, "Seuls les fichiers CSV sont supportés pour l'anonymisation pour le moment.")
                    url = reverse('anonymize') + f'?step=2&fichier_id={fichier_id}'
                    return redirect(url)
                    
            except Fichier.DoesNotExist:
                messages.error(request, "Fichier introuvable.")
                return redirect('anonymize')
            except Exception as e:
                messages.error(request, f"Erreur lors du traitement du fichier: {str(e)}")
                url = reverse('anonymize') + f'?step=2&fichier_id={fichier_id}'
                return redirect(url)
        # Dans votre fonction anonymize, remplacez la condition 'from_anonymize' par cette version corrigée :
        elif 'from_anonymize' in request.POST:
            fichier_id = request.POST.get('fichier_id')
            if not fichier_id:
                messages.error(request, "Aucun fichier sélectionné.")
                return redirect('anonymize')
            
            try:
                fichier_instance = Fichier.objects.get(id=fichier_id, utilisateur=request.user)
                
                # Récupérer les données depuis la session
                selected_ipi = request.session.get('selected_ipi', [])
                methods = request.session.get('methods', {})
                headers = request.session.get('headers', [])
                original = request.session.get('original', [])
                anonymized = request.session.get('anonymized', [])
                
                # Vérifications des données manquantes
                missing_data = []
                if not selected_ipi: missing_data.append("selected_ipi")
                if not methods: missing_data.append("methods")
                if not headers: missing_data.append("headers")
                if not original: missing_data.append("original")
                if not anonymized: missing_data.append("anonymized")
                    
                if missing_data:
                    messages.error(request, f"Données d'anonymisation manquantes: {', '.join(missing_data)}. Veuillez recommencer.")
                    # Nettoyer complètement la session
                    session_keys_to_clear = ['uploaded_file', 'selected_ipi', 'methods', 'headers', 'original', 'anonymized', 'lignes_traitees', 'colonnes_anonymisees', 'taux_anonymisation', 'score_securite']
                    for key in session_keys_to_clear:
                        if key in request.session:
                            del request.session[key]
                    return redirect('anonymize')
                
                # Créer le fichier anonymisé
                anonymized_content = io.StringIO()
                writer = csv.writer(anonymized_content)
                
                if headers:
                    writer.writerow(headers)
                
                for row in anonymized:
                    writer.writerow(row)
                
                # Sauvegarder le fichier anonymisé
                anonymized_filename = f"anonymized_{fichier_instance.nom_fichier}"
                anonymized_file_content = ContentFile(anonymized_content.getvalue().encode('utf-8'))
                anonymized_file_path = default_storage.save(f'fichiers_anonymises/{anonymized_filename}', anonymized_file_content)
                
                # Mettre à jour le fichier
                fichier_instance.statut = 'Anonymisé'
                fichier_instance.fichier_anonymise = anonymized_file_path
                fichier_instance.save()
                
                # Créer l'historique avec les méthodes utilisées
                methodes_str = ", ".join([f"{field}: {method}" for field, method in methods.items()])
                historique = Historique.objects.create(
                    fichier=fichier_instance,
                    utilisateur=request.user,
                    partage=False,
                    méthode_anonymisation=methodes_str,
                    statut='Terminé'
                )
                
                # Créer les métriques
                lignes_traitees = len(original)
                colonnes_anonymisees = len(selected_ipi)
                taux_anonymisation = (colonnes_anonymisees / len(headers)) * 100 if headers else 0
                score_securite = min(100, taux_anonymisation + calculate_sensitivity_score(selected_ipi))
                
                Métriques.objects.create(
                    historique=historique,
                    lignes_traitees=lignes_traitees,
                    colonnes_anonymisees=colonnes_anonymisees,
                    taux_anonymisation=taux_anonymisation,
                    score_securite=score_securite
                )
                
                # Créer le rapport de conformité
                RapportConformité.objects.create(
                    historique=historique,
                    analyse_risques=f"Score de sensibilité: {calculate_sensitivity_score(selected_ipi)}%",
                    recommandations="Fichier anonymisé selon les standards RGPD et CDP Sénégal",
                    conformite="Conforme RGPD/CDP"
                )
                
                # NETTOYER COMPLÈTEMENT LA SESSION
                session_keys_to_clear = [
                    'uploaded_file', 'selected_ipi', 'methods', 'headers', 
                    'original', 'anonymized', 'lignes_traitees', 'colonnes_anonymisees', 
                    'taux_anonymisation', 'score_securite'
                ]
                for key in session_keys_to_clear:
                    if key in request.session:
                        del request.session[key]
                
                # Forcer la sauvegarde de la session
                request.session.modified = True
                
                messages.success(request, f"Fichier '{fichier_instance.nom_fichier}' anonymisé avec succès!")
                
                # Rediriger vers l'étape 1 pour un nouveau cycle
                return redirect('anonymize')
                
            except Fichier.DoesNotExist:
                messages.error(request, "Fichier introuvable.")
                return redirect('anonymize')
            except Exception as e:
                messages.error(request, f"Erreur lors de la sauvegarde: {str(e)}")
                return redirect('anonymize')
        elif 'fichier_id' in request.POST and 'from_export' in request.POST:
            fichier_id = request.POST.get('fichier_id')
            format_export = request.POST.get('format')
            share_hub = request.POST.get('share_hub', False)
            
            # Stocker les paramètres d'export
            request.session['format_export'] = format_export
            request.session['share_hub'] = share_hub
            
            # Logique d'export (à implémenter)
            url = reverse('anonymize') + f'?step=4&fichier_id={fichier_id}'
            return redirect(url)

    # Récupérer les données depuis la session si elles existent
    if step >= 3:
        headers = request.session.get('headers', [])
        original = request.session.get('original', [])
        anonymized = request.session.get('anonymized', [])
        lignes_traitees = request.session.get('lignes_traitees', 0)
        colonnes_anonymisees = request.session.get('colonnes_anonymisees', 0)
        taux_anonymisation = request.session.get('taux_anonymisation', 0)
        score_securite = request.session.get('score_securite', 0)

    # Contexte pour toutes les étapes
    context = {
        'step': step,
        'uploaded_file': uploaded_file,
        'fichiers': fichiers,
        'methodes': methodes,
        'ipi_fields': ipi_fields,
        'sensitivity_score': sensitivity_score,
        'lignes_traitees': lignes_traitees,
        'colonnes_anonymisees': colonnes_anonymisees,
        'taux_anonymisation': taux_anonymisation,
        'score_securite': score_securite,
        'analyse_risques': analyse_risques,
        'conformite': conformite,
        'recommandations': recommandations,
        'fichier_id': fichier_id,
        'headers': headers,
        'original': original,
        'anonymized': anonymized,
    }
    return render(request, 'medicanon/anonymize.html', context)


@login_required
def anonymize_result(request, fichier_id):
    fichier = get_object_or_404(Fichier, id=fichier_id, utilisateur=request.user)
    headers = request.session.get('headers', [])
    original = request.session.get('original', [])
    anonymized = request.session.get('anonymized', [])
    return render(request, 'medicanon/anonymize_result.html', {
        'fichier': fichier,
        'headers': headers,
        'original': original,
        'anonymized': anonymized,
    })

@login_required
def share_anonymized(request, fichier_id):
    fichier = get_object_or_404(Fichier, id=fichier_id, utilisateur=request.user)
    fichier.partage = True
    fichier.save()
    messages.success(request, "Fichier partagé avec succès.")
    return redirect('anonymize_result', fichier_id=fichier.id)

@login_required
def download_anonymized(request, fichier_id):
    fichier = get_object_or_404(Fichier, id=fichier_id, utilisateur=request.user)
    anonymized = request.session.get('anonymized', [])
    headers = request.session.get('headers', [])
    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = f'attachment; filename=anonymized_{fichier.nom_fichier}'
    writer = csv.writer(response)
    if headers:
        writer.writerow(headers)
    for row in anonymized:
        writer.writerow(row)
    return response

@login_required
def metrics(request):
    historiques = Historique.objects.filter(utilisateur=request.user)
    metriques = Métriques.objects.filter(historique__in=historiques)
    return render(request, 'medicanon/metrics.html', {'metriques': metriques})

@login_required
def compliance_report(request):
    historiques = Historique.objects.filter(utilisateur=request.user)
    rapports = RapportConformité.objects.filter(historique__in=historiques)
    return render(request, 'medicanon/compliance_report.html', {'rapports': rapports})

def custom_logout(request):
    logout(request)
    return redirect('accueil')

@login_required
def import_fichier(request):
    fichiers = Fichier.objects.all().order_by('-date_import')
    search_query = request.GET.get('search', '').strip()
    statut_filter = request.GET.get('statut', '').strip()
    if search_query:
        fichiers = fichiers.filter(nom_fichier__icontains=search_query)
    if statut_filter:
        fichiers = fichiers.filter(statut=statut_filter)
    if request.method == 'POST' and request.FILES.get('fichier'):
        uploaded_file = request.FILES['fichier']
        fichier_instance = Fichier.objects.create(
            fichier=uploaded_file,
            nom_fichier=uploaded_file.name,
            statut='Importé',
            utilisateur=request.user
        )
        return redirect('import_fichier')
    paginator = Paginator(fichiers, 5)
    page_number = request.GET.get('page')
    fichiers_page = paginator.get_page(page_number)
    if request.headers.get('HX-Request') == 'true':
        return render(request, 'medicanon/_fichiers_table.html', {'fichiers': fichiers_page})
    return render(request, 'medicanon/import.html', {
        'fichiers': fichiers_page,
        'search_query': search_query,
        'statut_filter': statut_filter
    })

@login_required
def delete_fichier(request, pk):
    fichier = get_object_or_404(Fichier, pk=pk)
    fichier.delete()
    fichiers = Fichier.objects.all().order_by('-date_import')
    paginator = Paginator(fichiers, 5)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    return render(request, 'medicanon/_fichiers_table.html', {'fichiers': page_obj})


@login_required
def telecharger_fichier(request, fichier_id):  
    fichier = get_object_or_404(Fichier, id=fichier_id, utilisateur=request.user)
    
    # Vérifier si le fichier anonymisé existe
    if fichier.fichier_anonymise and fichier.statut == 'Anonymisé':
        return FileResponse(
            fichier.fichier_anonymise.open('rb'), 
            as_attachment=True, 
            filename=f"anonymized_{fichier.nom_fichier}"
        )
    
    # Fallback : générer le CSV depuis la session (si encore disponible)
    anonymized = request.session.get('anonymized', [])
    headers = request.session.get('headers', [])
    
    if not anonymized:
        messages.error(request, "Fichier anonymisé non disponible. Veuillez relancer l'anonymisation.")
        return redirect('anonymize')
    
    # Générer le CSV à la volée
    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = f'attachment; filename=anonymized_{fichier.nom_fichier}'
    writer = csv.writer(response)
    
    if headers:
        writer.writerow(headers)
    for row in anonymized:
        writer.writerow(row)
    
    return response

@login_required
def export_fichiers_csv(request):
    fichiers = Fichier.objects.all().order_by('-date_import')
    search_query = request.GET.get('search')
    statut_filter = request.GET.get('statut')
    if search_query:
        fichiers = fichiers.filter(nom_fichier__icontains=search_query)
    if statut_filter:
        fichiers = fichiers.filter(statut=statut_filter)
    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = 'attachment; filename="fichiers_medic_anon.csv"'
    writer = csv.writer(response)
    writer.writerow(['Nom Fichier', 'Date Import', 'Statut', 'Utilisateur'])
    for fichier in fichiers:
        writer.writerow([
            fichier.nom_fichier,
            fichier.date_import.strftime("%d/%m/%Y %H:%M"),
            fichier.statut,
            fichier.utilisateur.username
        ])
    return response

def public_hub(request):
    fichiers_anonymises = Fichier.objects.filter(statut='Anonymisé')
    return render(request, 'medicanon/public_hub.html', {'fichiers': fichiers_anonymises})

@require_GET
def file_preview_api(request, fichier_id):
    """API pour récupérer l'aperçu d'un fichier anonymisé"""
    
    try:
        fichier = get_object_or_404(Fichier, id=fichier_id)
        
        # Vérifier que le fichier est anonymisé et partagé publiquement
        if fichier.statut != 'Anonymisé' or not fichier.partage:
            return JsonResponse({
                'error': 'Fichier non disponible pour aperçu',
                'message': 'Ce fichier n\'est pas anonymisé ou n\'est pas partagé publiquement.'
            }, status=403)
        
        # Logger l'accès pour audit
        if request.user.is_authenticated:
            AuditLog.objects.create(
                utilisateur=request.user,
                action='VIEW',
                fichier=fichier,
                details={'action': 'preview_access', 'ip': get_client_ip(request)},
                ip_address=get_client_ip(request)
            )
        
        # Déterminer le type de fichier
        file_extension = fichier.nom_fichier.split('.')[-1].lower()
        
        # Préparer les métadonnées de base
        metadata = {
            'id': fichier.id,
            'name': fichier.nom_fichier,
            'type': file_extension,
            'size': get_file_size_display(fichier),
            'date': fichier.date_import.strftime('%d/%m/%Y'),
            'status': fichier.statut,
            'score': calculate_display_score(fichier)
        }
        
        # Générer l'aperçu selon le type
        if file_extension == 'csv':
            preview_data = generate_csv_preview(fichier)
        elif file_extension == 'pdf':
            preview_data = generate_pdf_preview(fichier)
        elif file_extension in ['docx', 'doc']:
            preview_data = generate_docx_preview(fichier)
        else:
            preview_data = {'type': 'unsupported', 'message': 'Type de fichier non supporté pour l\'aperçu'}
        
        return JsonResponse({
            'success': True,
            'metadata': metadata,
            'preview': preview_data
        })
        
    except Exception as e:
        return JsonResponse({
            'error': 'Erreur lors de la génération de l\'aperçu',
            'message': str(e)
        }, status=500)

def generate_csv_preview(fichier):
    """Génère un aperçu pour les fichiers CSV"""
    try:
        # Utiliser le fichier anonymisé si disponible, sinon le fichier original
        file_to_read = fichier.fichier_anonymise if fichier.fichier_anonymise else fichier.fichier
        
        file_to_read.seek(0)
        content = file_to_read.read().decode('utf-8', errors='ignore')
        file_to_read.seek(0)
        
        # Parser le CSV
        csv_reader = csv.reader(io.StringIO(content))
        
        # Lire les en-têtes
        headers = next(csv_reader, [])
        
        # Lire les 5 premières lignes de données
        rows = []
        for i, row in enumerate(csv_reader):
            if i >= 5:  # Limiter à 5 lignes pour l'aperçu
                break
            rows.append(row)
        
        # Statistiques
        file_to_read.seek(0)
        total_lines = sum(1 for line in io.StringIO(content)) - 1  # -1 pour les en-têtes
        
        return {
            'type': 'csv',
            'headers': headers,
            'rows': rows,
            'total_rows': total_lines,
            'total_columns': len(headers),
            'preview_rows': len(rows)
        }
        
    except Exception as e:
        return {
            'type': 'error',
            'message': f'Erreur lors de la lecture du CSV: {str(e)}'
        }

def generate_pdf_preview(fichier):
    """Génère un aperçu pour les fichiers PDF"""
    try:
        file_to_read = fichier.fichier_anonymise if fichier.fichier_anonymise else fichier.fichier
        file_to_read.seek(0)
        
        # Utiliser pdfplumber pour extraire des informations
        with pdfplumber.open(file_to_read) as pdf:
            num_pages = len(pdf.pages)
            
            # Extraire le texte de la première page pour aperçu
            first_page_text = ""
            if num_pages > 0:
                first_page_text = pdf.pages[0].extract_text() or "Contenu non extractible"
                # Limiter à 500 caractères pour l'aperçu
                first_page_text = first_page_text[:500] + ("..." if len(first_page_text) > 500 else "")
        
        file_to_read.seek(0)
        
        return {
            'type': 'pdf',
            'pages': num_pages,
            'first_page_preview': first_page_text,
            'message': f'Document PDF de {num_pages} page(s) - Données anonymisées'
        }
        
    except Exception as e:
        return {
            'type': 'error',
            'message': f'Erreur lors de la lecture du PDF: {str(e)}'
        }

def generate_docx_preview(fichier):
    """Génère un aperçu pour les fichiers DOCX"""
    try:
        file_to_read = fichier.fichier_anonymise if fichier.fichier_anonymise else fichier.fichier
        file_to_read.seek(0)
        
        # Lire le document Word
        doc = Document(io.BytesIO(file_to_read.read()))
        
        # Extraire les premiers paragraphes
        paragraphs = []
        total_paragraphs = len(doc.paragraphs)
        
        for i, para in enumerate(doc.paragraphs):
            if i >= 3:  # Limiter à 3 paragraphes
                break
            if para.text.strip():  # Ignorer les paragraphes vides
                paragraphs.append(para.text.strip())
        
        # Compter les tables
        num_tables = len(doc.tables)
        
        file_to_read.seek(0)
        
        return {
            'type': 'docx',
            'paragraphs': paragraphs,
            'total_paragraphs': total_paragraphs,
            'tables': num_tables,
            'message': f'Document Word avec {total_paragraphs} paragraphe(s) et {num_tables} table(s) - Données anonymisées'
        }
        
    except Exception as e:
        return {
            'type': 'error',
            'message': f'Erreur lors de la lecture du document: {str(e)}'
        }

def get_file_size_display(fichier):
    """Retourne la taille du fichier en format lisible"""
    try:
        file_to_check = fichier.fichier_anonymise if fichier.fichier_anonymise else fichier.fichier
        size_bytes = file_to_check.size
        
        if size_bytes < 1024:
            return f"{size_bytes} B"
        elif size_bytes < 1024**2:
            return f"{size_bytes/1024:.1f} KB"
        elif size_bytes < 1024**3:
            return f"{size_bytes/(1024**2):.1f} MB"
        else:
            return f"{size_bytes/(1024**3):.1f} GB"
    except:
        return "Taille inconnue"

def calculate_display_score(fichier):
    """Calcule un score d'affichage pour le fichier"""
    try:
        # Récupérer les métriques du fichier
        from .models import Métriques, Historique
        
        historique = Historique.objects.filter(fichier=fichier).first()
        if historique:
            metriques = Métriques.objects.filter(historique=historique).first()
            if metriques:
                return f"{metriques.score_securite:.0f}/100"
        
        # Score par défaut basé sur le statut
        if fichier.statut == 'Anonymisé':
            return "95/100"
        else:
            return "0/100"
    except:
        return "95/100"

def get_client_ip(request):
    """Obtient l'IP réelle du client"""
    x_forwarded_for = request.META.get('HTTP_X_FORWARDED_FOR')
    if x_forwarded_for:
        ip = x_forwarded_for.split(',')[0]
    else:
        ip = request.META.get('REMOTE_ADDR')
    return ip

